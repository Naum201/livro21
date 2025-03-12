from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from werkzeug.utils import secure_filename
import re
from app import app, db
from datetime import datetime
import os
from docx import Document
from app import Usuario, Seguidores, Obra, Capitulo, Comentario, ObraTag, Tag, Voto, Genero, Postagem, Notificacao, ObraPlataforma, Plataforma
from difflib import SequenceMatcher
from sqlalchemy import func

UPLOAD_FOLDER = 'static/uploads/capas'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# Novo diretório para imagens de capítulos
CAPITULO_UPLOAD_FOLDER = 'static/uploads/capitulos'
app.config['CAPITULO_UPLOAD_FOLDER'] = CAPITULO_UPLOAD_FOLDER
app.config['CAPITULO_UPLOAD_FOLDER'] = CAPITULO_UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/')
def index():
    obras_recomendadas = Obra.query.order_by(Obra.id.desc()).limit(4).all()
    return render_template('index.html', obras_recomendadas=obras_recomendadas)

@app.route("/explorar")
def explorar():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    usuario_id = session['user_id']
    usuario = Usuario.query.get(usuario_id)
     # Obtendo o usuário logado
    usuario_logado_id = session.get('user_id')
    if usuario_logado_id:
        usuario_logado = Usuario.query.get_or_404(usuario_logado_id)
    else:
        usuario_logado = None  # Caso não exista um usuário logado
    
    filtro = request.args.get("filtro", "todas")  # Pega o filtro da URL (padrão: todas)
    
    if filtro == "populares":
        obras = Obra.query.order_by(Obra.likes.desc()).all()  # Exemplo: ordenando por curtidas
    elif filtro == "recentes":
        obras = Obra.query.order_by(Obra.data_criacao.desc()).all()  # Ordena por data
    else:
        obras = Obra.query.all()  # Retorna todas as obras

    return render_template("usuario/explorar.html", obras=obras, usuario=usuario, usuario_logado=usuario_logado)  # Passa o usuário para o template

@app.route("/explorar/genero/<int:genero_id>")
def explorar_por_genero(genero_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    usuario_id = session['user_id']
    usuario = Usuario.query.get(usuario_id)
     # Obtendo o usuário logado
    usuario_logado_id = session.get('user_id')
    if usuario_logado_id:
        usuario_logado = Usuario.query.get_or_404(usuario_logado_id)
    else:
        usuario_logado = None  # Caso não exista um usuário logado

    # Obtemos o gênero selecionado ou levanta um erro 404 se não encontrado
    genero = Genero.query.get_or_404(genero_id)

    # Filtra as obras pelo gênero
    obras = Obra.query.filter_by(genero_id=genero.id).all()

    return render_template("usuario/explorar_por_genero.html", obras=obras, genero=genero, usuario=usuario, usuario_logado=usuario_logado)

@app.route('/pesquisar', methods=['GET'])
def pesquisar():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    usuario_id = session['user_id']
    usuario = Usuario.query.get(usuario_id)
     # Obtendo o usuário logado
    usuario_logado_id = session.get('user_id')
    if usuario_logado_id:
        usuario_logado = Usuario.query.get_or_404(usuario_logado_id)
    else:
        usuario_logado = None  # Caso não exista um usuário logado
    
    # Obtém o termo de pesquisa da URL ou do formulário
    query = request.args.get('query', '')  # Assumindo que a query seja passada como parâmetro de URL
    
    # Realiza a busca nas obras (livros)
    obras = Obra.query.filter(
        func.lower(Obra.titulo).like(f'%{query.lower()}%') |
        func.lower(Obra.sinopse).like(f'%{query.lower()}%')  # Alterado de 'conteudo' para 'sinopse'
    ).all()

    # Realiza a busca nos usuários
    usuarios = Usuario.query.filter(
        func.lower(Usuario.username).like(f'%{query.lower()}%') |
        func.lower(Usuario.username).like(f'%{query.lower()}%')
    ).all()

    # Renderiza o template de resultados com os dados encontrados e o usuário logado
    return render_template('usuario/pesquisa_resultados.html', query=query, obras=obras, usuarios=usuarios, usuario=usuario, usuario_logado=usuario_logado)

def get_user_logged_in():
    # Supondo que você armazene o ID do usuário na sessão
    user_id = session.get('user_id')
    if user_id:
        return Usuario.query.get(user_id)
    return None

@app.route('/seguir/<int:user_id>', methods=['POST'])
def seguir(user_id):
    usuario_logado = get_user_logged_in()
    
    if not usuario_logado:
        return jsonify({"status": "error", "message": "Usuário não autenticado"}), 403

    if usuario_logado.id == user_id:
        return jsonify({"status": "error", "message": "Você não pode seguir a si mesmo"}), 400

    if Seguidores.query.filter_by(seguidor_id=usuario_logado.id, seguido_id=user_id).first():
        return jsonify({"status": "error", "message": "Você já segue esse usuário"}), 400

    novo_seguimento = Seguidores(seguidor_id=usuario_logado.id, seguido_id=user_id)
    db.session.add(novo_seguimento)
    db.session.commit()

    return jsonify({"status": "success", "message": "Agora você segue esse usuário"})


@app.route('/deixar_de_seguir/<int:user_id>', methods=['POST'])
def deixar_de_seguir(user_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    # Verificar se o usuário está realmente seguindo
    usuario_seguido = Usuario.query.get(user_id)
    if not usuario_seguido:
        flash('Usuário não encontrado.', 'danger')
        return redirect(url_for('pesquisar'))

    # Evitar que o usuário deixe de seguir a si mesmo
    if user_id == session['user_id']:
        flash('Você não pode deixar de seguir a si mesmo.', 'danger')
        return redirect(url_for('pesquisar'))

    # Remover o seguimento
    seguimento = Seguidores.query.filter_by(usuario_id=session['user_id'], usuario_seguido_id=user_id).first()
    if seguimento:
        db.session.delete(seguimento)
        db.session.commit()

        flash('Você deixou de seguir esse usuário!', 'success')
    else:
        flash('Você não segue este usuário.', 'danger')

    return redirect(url_for('pesquisar'))


def verificar_plagio(conteudo, lista_existente):
    for existente in lista_existente:
        similaridade = SequenceMatcher(None, conteudo, existente).ratio()
        if similaridade > 0.8:  # Limite de 80% de similaridade
            return True, existente
    return False, None

@app.route('/criar_obras', methods=['GET', 'POST'])
def criar_obras():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    usuario_logado_id = session.get('user_id')
    usuario_logado = Usuario.query.get(usuario_logado_id) if usuario_logado_id else None

    usuario_id = session['user_id']
    usuario = Usuario.query.get(usuario_id)
    
    generos = Genero.query.all()
    plataformas = Plataforma.query.all()
    obras_publicadas = Obra.query.filter_by(usuario_id=usuario.id, status='Publicado').all()
    obras_rascunhos = Obra.query.filter_by(usuario_id=usuario.id, status='Rascunho').all()

    if request.method == 'POST':
        titulo = request.form['titulo']
        sinopse = request.form['sinopse']
        genero_id = request.form['genero']
        classificacao = request.form['classificacao']
        etapa = request.form['etapa']
        status = request.form['status']
        plataformas_ids = request.form.getlist('plataformas')
        tags_input = request.form['tags']  # Obter as tags do formulário

        # Processar as tags, separando por vírgula e limpando espaços
        tags = [tag.strip() for tag in tags_input.split(',') if tag.strip()]

        # Criar novas tags se não existirem
        tags_obj = []
        for tag_nome in tags:
            tag = Tag.query.filter_by(nome=tag_nome).first()
            if not tag:
                tag = Tag(nome=tag_nome)  # Criar nova tag se não existir
                db.session.add(tag)
            tags_obj.append(tag)

        capa = None
        if 'capa' in request.files:
            capa_file = request.files['capa']
            if capa_file.filename != '':
                capa = capa_file.filename
                capa_file.save(f'static/uploads/capas/{capa}')

        # Obter sinopses existentes no banco de dados
        obras_existentes = [obra.sinopse for obra in Obra.query.all()]
        
        # Verificar plágio no banco de dados
        plagio_detectado, obra_similar = verificar_plagio(sinopse, obras_existentes)
        if plagio_detectado:
            flash(f'Plágio detectado! O conteúdo é semelhante à obra com sinopse: "{obra_similar}".', 'danger')
            return redirect(url_for('criar_obras'))

        # Se não houve plágio, criar a nova obra
        nova_obra = Obra(
            titulo=titulo,
            sinopse=sinopse,
            genero_id=genero_id,
            classificacao=classificacao,
            etapa=etapa,
            status=status,
            capa=capa,
            usuario_id=usuario.id,
        )
        db.session.add(nova_obra)
        db.session.commit()

        # Adicionar plataformas associadas à obra
        for plataforma_id in plataformas_ids:
            obra_plataforma = ObraPlataforma(
                obra_id=nova_obra.id,
                plataforma_id=plataforma_id
            )
            db.session.add(obra_plataforma)

        # Associar as tags à obra
        for tag in tags_obj:
            nova_obra.tags.append(tag)

        db.session.commit()
        flash('Obra criada com sucesso!', 'success')
        return redirect(url_for('criar_capitulos', obra_id=nova_obra.id))

    return render_template(
        'usuario/criar_obras.html',
        generos=generos,
        plataformas=plataformas,
        usuario=usuario,
        obras_publicadas=obras_publicadas,
        obras_rascunhos=obras_rascunhos,
        usuario_logado=usuario_logado
    )
@app.route('/tags')
def tags():
    # Obtendo todas as tags do banco de dados
    tags = Tag.query.all()

    # Para cada tag, calcular as obras mais lidas relacionadas a ela
    obras_por_tag = {}
    for tag in tags:
        obras = db.session.query(Obra).join(ObraTag).filter(ObraTag.tag_id == tag.id).order_by(Obra.visualizacoes.desc()).limit(5).all()
        obras_por_tag[tag] = obras

    return render_template('tags.html', tags=tags, obras_por_tag=obras_por_tag)
# Supondo que você já tenha configurado o Flask e o SQLAlchemy
@app.route('/minhas_obras')
def minhas_obras():
    usuario_logado_id = session.get('user_id')
    usuario_logado = Usuario.query.get(usuario_logado_id) if usuario_logado_id else None
    
    # Consultando as obras do usuário
    obras = Obra.query.filter_by(usuario_id=usuario_logado_id).all()
    
    obras_info = []
    for obra in obras:
        # Contando capítulos e comentários (assumindo que os comentários estão relacionados com os capítulos)
        quantidade_capitulos = len(obra.capitulos)
        quantidade_votos = obra.votos
        quantidade_comentarios = sum(len(capitulo.comentarios) for capitulo in obra.capitulos)  # Comentários de todos os capítulos

        # Montando as informações da obra
        obras_info.append({
            'capa': obra.capa,
            'titulo': obra.titulo,
            'classificacao': obra.classificacao,
            'sinopse': obra.sinopse,
            'quantidade_capitulos': quantidade_capitulos,
            'quantidade_votos': quantidade_votos,
            'quantidade_comentarios': quantidade_comentarios,
        })
    
    return render_template('usuario/minhas_obras.html', obras=obras_info, usuario_logado=usuario_logado)



@app.route('/obra/<int:obra_id>')
def obra_detalhes(obra_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    usuario_id = session['user_id']
    usuario = Usuario.query.get(usuario_id)
     # Obtendo o usuário logado
    usuario_logado_id = session.get('user_id')
    if usuario_logado_id:
        usuario_logado = Usuario.query.get_or_404(usuario_logado_id)
    else:
        usuario_logado = None  # Caso não exista um usuário logado
    
    # Buscar a obra pelo ID
    obra = Obra.query.get_or_404(obra_id)

    # Incrementar as visualizações
    if usuario_id != obra.usuario_id:  # Se o usuário não for o autor
        # Contabilizar a visualização, caso não seja o autor da obra
        obra.visualizacoes += 1
        db.session.commit()

    # Contar comentários e votos
    total_comentarios = Comentario.query.filter(Comentario.capitulo_id.in_([capitulo.id for capitulo in obra.capitulos])).count()
    total_votos = obra.votos

    # Obter capítulos publicados da obra
    capitulos_publicados = Capitulo.query.filter_by(obra_id=obra.id, status='Publicado').all()

    return render_template(
        'usuario/obra_detalhes.html',
        obra=obra,
        total_comentarios=total_comentarios,
        total_votos=total_votos,
        capitulos_publicados=capitulos_publicados,  # Passando os capítulos publicados para o template
        usuario=usuario,
        usuario_logado=usuario_logado
    )

# Função para extrair conteúdo de arquivos .docx
def extract_text_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Função para extrair conteúdo de arquivos .docx, incluindo tabelas
def extract_text_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    full_html = []

    for para in doc.paragraphs:
        # Verificar se o parágrafo é formatado (negrito, itálico, etc.)
        para_html = ""
        for run in para.runs:
            if run.bold:
                para_html += f"<strong>{run.text}</strong>"
            elif run.italic:
                para_html += f"<em>{run.text}</em>"
            elif run.underline:
                para_html += f"<u>{run.text}</u>"
            else:
                para_html += run.text
        full_html.append(f"<p>{para_html}</p>")  # Envolvendo o texto em <p> para manter a estrutura de parágrafo

    # Gerar HTML para as tabelas
    for table in doc.tables:
        table_html = '<table class="table table-bordered">'
        for row in table.rows:
            table_html += '<tr>'
            for cell in row.cells:
                table_html += f'<td>{cell.text}</td>'
            table_html += '</tr>'
        table_html += '</table>'
        full_html.append(table_html)

    return ''.join(full_html)

@app.route('/criar_capitulos/<int:obra_id>', methods=['GET', 'POST'])
def criar_capitulos(obra_id):
    obra = Obra.query.get_or_404(obra_id)

    if request.method == 'POST':
        try:
            titulo_capitulo = request.form['titulo']
            status = request.form['status']
            arquivo = request.files.get('arquivo')
            conteudo_capitulo = None

            if arquivo and allowed_file(arquivo.filename):
                filename = secure_filename(arquivo.filename)
                file_path = os.path.join(app.config['CAPITULO_UPLOAD_FOLDER'], filename)
                arquivo.save(file_path)
                conteudo_capitulo = extract_text_from_docx(file_path)
            elif 'conteudo' in request.form:
                conteudo_capitulo = request.form['conteudo']

            if not conteudo_capitulo:
                flash("Por favor, forneça o conteúdo do capítulo ou envie um arquivo DOCX.", "danger")
                return redirect(url_for('criar_capitulos', obra_id=obra.id))

            # Verificar plágio nos capítulos existentes
            capitulos_existentes = [capitulo.conteudo for capitulo in Capitulo.query.all()]
            plagio_detectado, capitulo_similar = verificar_plagio(conteudo_capitulo, capitulos_existentes)
            if plagio_detectado:
                flash(f'Plágio detectado! O conteúdo é semelhante ao capítulo "{capitulo_similar}".', 'danger')
                return redirect(url_for('criar_capitulos', obra_id=obra.id))

            novo_capitulo = Capitulo(
                titulo=titulo_capitulo,
                conteudo=conteudo_capitulo,
                status=status,
                obra_id=obra.id
            )

            db.session.add(novo_capitulo)
            db.session.commit()

            flash("Capítulo criado com sucesso!", "success")

            # 1. Enviar notificações para os seguidores da obra
            seguidores = obra.autor.seguindo.all() if obra.autor else []
            for seguidor in seguidores:
                notificacao = Notificacao(
                    mensagem=f"Novo capítulo '{novo_capitulo.titulo}' na obra '{obra.titulo}' que você está seguindo!",
                    usuario_id=seguidor.id,
                    data_criacao=datetime.utcnow()
                )
                db.session.add(notificacao)

            # 2. Enviar notificações para os leitores do projeto de leitura
            if hasattr(obra, 'projetos') and obra.projetos:
                for projeto in obra.projetos:
                    for usuario in projeto.usuarios:
                        notificacao = Notificacao(
                            mensagem=f"Novo capítulo '{novo_capitulo.titulo}' na obra '{obra.titulo}' do projeto de leitura '{projeto.nome_projeto}' que você está participando!",
                            usuario_id=usuario.id,
                            data_criacao=datetime.utcnow()
                        )
                        db.session.add(notificacao)

            db.session.commit()
            return redirect(url_for('visualizacao_capitulo', capitulo_id=novo_capitulo.id))

        except Exception as e:  # Adicionando um bloco except
            flash(f"Ocorreu um erro ao criar o capítulo: {str(e)}", "danger")
            db.session.rollback()
            return redirect(url_for('criar_capitulos', obra_id=obra.id))

    return render_template('usuario/criar_capitulos.html', obra=obra)

# Função para carregar as imagens e retornar a URL
@app.route('/upload_imagem/<int:obra_id>', methods=['POST'])
def upload_imagem(obra_id):
    if 'image' not in request.files:
        return jsonify({'error': 'Nenhuma imagem foi enviada'}), 400
    
    file = request.files['image']
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['CAPITULO_UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Criar URL para acessar a imagem
        image_url = url_for('static', filename=f'uploads/capitulos/{filename}')
        
        # Aqui você pode salvar a URL da imagem diretamente no campo 'imagem' do Capítulo
        return jsonify({'url': image_url}), 200
    else:
        return jsonify({'error': 'Arquivo não permitido'}), 400

    
@app.route('/editar_obra/<int:obra_id>', methods=['GET', 'POST'])
def editar_obra(obra_id):
    obra = Obra.query.get_or_404(obra_id)  # Fetch the work by its ID
    generos = Genero.query.all()  # Fetch all genres
    plataformas = Plataforma.query.all()  # Fetch all platforms

    if request.method == 'POST':
        # Get form data
        titulo = request.form['titulo']
        sinopse = request.form['sinopse']
        genero_id = request.form['genero']
        classificacao = request.form['classificacao']
        tag = request.form['tag']
        etapa = request.form['etapa']
        status = request.form['status']
        plataforma_id = request.form['plataformas']

        # Handle file upload for the cover image
        capa = request.files.get('capa')
        if capa:
            capa_filename = secure_filename(capa.filename)
            capa.save(os.path.join(app.config['UPLOAD_FOLDER'], capa_filename))
        else:
            capa_filename = obra.capa  # Keep the current cover if no new file is uploaded

        # Update the work details
        obra.titulo = titulo
        obra.sinopse = sinopse
        obra.genero_id = genero_id
        obra.classificacao = classificacao
        obra.tag = tag
        obra.etapa = etapa
        obra.status = status
        obra.plataforma_id = plataforma_id
        obra.capa = capa_filename  # Save the new or existing cover image

        db.session.commit()  # Save changes to the database

        return redirect(url_for('obra_detalhes', obra_id=obra.id))  # Redirect to the work page

    return render_template('usuario/editar_obra.html', obra=obra, generos=generos, plataformas=plataformas)

@app.route('/editar_capitulo/<int:capitulo_id>', methods=['GET', 'POST'])
def editar_capitulo(capitulo_id):
    # Retrieve the capitulo (chapter) from the database
    capitulo = Capitulo.query.get_or_404(capitulo_id)
    
    # Retrieve the associated obra (work) from the capitulo
    obra = capitulo.obra  # 'obra' is a relationship defined in the Capitulo model
    
    # Remove HTML tags from capitulo.conteudo for display
    capitulo_conteudo_stripped = re.sub(r'<.*?>', '', capitulo.conteudo)

    if request.method == 'POST':
        # Use .get() to avoid the KeyError if 'conteudo' is missing
        capitulo.titulo = request.form.get('titulo')
        capitulo.conteudo = request.form.get('conteudo')  # Using .get() to avoid KeyError
        capitulo.status = request.form.get('status')

        # Commit to save changes in the database
        db.session.commit()

        flash("Capítulo editado com sucesso!", "success")

        # Send notifications (continue with your original logic)
        seguidores = obra.autor.seguindo.all()
        for seguidor in seguidores:
            notificacao = Notificacao(
                mensagem=f"O capítulo '{capitulo.titulo}' foi editado na obra '{obra.titulo}' que você está seguindo!",
                usuario_id=seguidor.id,
                data_criacao=datetime.utcnow()
            )
            db.session.add(notificacao)

        if obra.projetos:
            for projeto in obra.projetos:
                for usuario in projeto.usuarios:
                    notificacao = Notificacao(
                        mensagem=f"O capítulo '{capitulo.titulo}' foi editado na obra '{obra.titulo}' do projeto de leitura '{projeto.nome_projeto}' que você está participando!",
                        usuario_id=usuario.id,
                        data_criacao=datetime.utcnow()
                    )
                    db.session.add(notificacao)

        db.session.commit()

        # Redirect to the chapter's visualização page
        return redirect(url_for('visualizacao_capitulo', capitulo_id=capitulo.id))

    # Pass both 'capitulo', 'obra', and 'capitulo_conteudo_stripped' to the template
    return render_template('usuario/editar_capitulo.html', capitulo=capitulo, obra=obra, capitulo_conteudo_stripped=capitulo_conteudo_stripped)

def criar_notificacao(acao, usuario_que_realizou_acao, usuario_seguindo):
    mensagem = ""
    
    if acao == "Postagem":
        mensagem = f"{usuario_que_realizou_acao.username} publicou uma nova postagem!"
    elif acao == "Novo Capítulo":
        mensagem = f"{usuario_que_realizou_acao.username} publicou um novo capítulo!"
    
    # Criar a notificação para cada seguidor
    for seguidor in usuario_seguindo:
        notificacao = Notificacao(
            usuario_id=seguidor.id,  # Quem vai receber a notificação
            acao=acao,                # Tipo de ação
            mensagem=mensagem,        # Mensagem da notificação
        )
        db.session.add(notificacao)
    db.session.commit()
@app.route('/notificacoes')
def notificacoes():
    # Obtendo o usuário logado
    usuario_logado_id = session.get('user_id')
    usuario_logado = Usuario.query.get(usuario_logado_id) if usuario_logado_id else None
    
    if usuario_logado:
        # 1. Buscar capítulos recentes das obras que o usuário está lendo (obras que ele é autor ou está seguindo)
        obras_seguidas = usuario_logado.seguindo.all()
        obras_autoria = usuario_logado.obras
        capitulos_recentemente_publicados = []

        for obra in obras_seguidas:
            capitulos_recentemente_publicados += obra.capitulos[:1]  # Pegando os 5 últimos capítulos (ajuste conforme necessário)

        for obra in obras_autoria:
            capitulos_recentemente_publicados += obra.capitulos[:1]  # Pegando os 5 últimos capítulos (ajuste conforme necessário)

        # 2. Buscar postagens dos usuários que o usuário logado segue
        postagens_seguidas = Postagem.query.filter(Postagem.usuario_id.in_([seguidor.seguidor_id for seguidor in usuario_logado.seguindo])).all()

        # 3. Adicionar as notificações relacionadas a essas postagens e capítulos
        notificacoes = []

        # Adicionando notificações de capítulos
        for capitulo in capitulos_recentemente_publicados:
            # Verificando se a notificação já existe (para evitar duplicação)
            notificacao_existente = Notificacao.query.filter_by(
                mensagem=f"Novo capítulo: {capitulo.titulo} na obra '{capitulo.obra.titulo}'",
                usuario_id=usuario_logado.id
            ).first()
            if not notificacao_existente:
                notificacao = Notificacao(
                    mensagem=f"Novo capítulo: {capitulo.titulo} na obra '{capitulo.obra.titulo}'",
                    usuario_id=usuario_logado.id,
                    data_criacao=datetime.utcnow(),
                )
                notificacoes.append(notificacao)

        # Adicionando notificações de postagens
        for postagem in postagens_seguidas:
            # Verificando se a notificação já existe (para evitar duplicação)
            notificacao_existente = Notificacao.query.filter_by(
                mensagem=f"Nova postagem: {postagem.titulo} por {postagem.usuario.username}",
                usuario_id=usuario_logado.id
            ).first()
            if not notificacao_existente:
                notificacao = Notificacao(
                    mensagem=f"Nova postagem: {postagem.titulo} por {postagem.usuario.username}",
                    usuario_id=usuario_logado.id,
                    data_criacao=datetime.utcnow(),
                )
                notificacoes.append(notificacao)

        # Adicionando as notificações ao banco de dados
        db.session.add_all(notificacoes)
        db.session.commit()

        # Filtrar as notificações não vistas
        notificacoes_nao_vistas = Notificacao.query.filter_by(usuario_id=usuario_logado.id, visto=False).all()

        return render_template('usuario/notificacoes.html', notificacoes=notificacoes_nao_vistas, usuario_logado=usuario_logado)

@app.route('/marcar_como_visto/<int:id>', methods=['POST'])
def marcar_como_visto(id):
    usuario_logado_id = session.get('user_id')
    usuario_logado = Usuario.query.get_or_404(usuario_logado_id)

    # Procurando a notificação
    notificacao = Notificacao.query.filter_by(id=id, usuario_id=usuario_logado.id).first()

    if notificacao:
        notificacao.visto = True
        db.session.commit()
        return jsonify({"success": True})
    else:
        return jsonify({"success": False, "message": "Notificação não encontrada."}), 404


@app.route('/biblioteca')
def biblioteca():
    # Obtendo o usuário logado
    usuario_logado_id = session.get('user_id')
    usuario_logado = Usuario.query.get(usuario_logado_id) if usuario_logado_id else None
    
    if usuario_logado:
        # Obter as obras que o usuário está lendo:
        # Obras que ele segue
        obras_seguidas = [seguidor.obra for seguidor in usuario_logado.seguindo]
        
        # Obras que ele criou
        obras_autoria = usuario_logado.obras
        
        # Combinar as duas listas de obras (seguindo e autoria)
        obras_do_usuario = obras_seguidas + obras_autoria
        
        return render_template('usuario/biblioteca.html', obras=obras_do_usuario, usuario_logado=usuario_logado)
    
    else:
        return render_template('usuario/biblioteca.html', obras=[], usuario_logado=None)
